"""pypdf + python-docx — the `document_parse` pin's rung-0.5/1 seat. VENDOR
CODE LIVES HERE, per Law 1: `core` never imports either library.

    document_parse: {dev: local_pypdf_docx}
    — docs/architecture/plates/04-pin-out-map.md

PDF GETS REAL PAGE ANCHORS; DOCX GETS ONE LOGICAL PAGE. `pypdf.PdfReader`
exposes exact page boundaries the way a reader sees them, which is what
E16-06's "companion guide p.14" citation needs — the common case, since a
payer's companion guide arrives as a PDF. `python-docx` has no equivalent: a
`.docx` records paragraphs and tables, not print-time page breaks, and
fabricating page numbers from a heuristic would be inventing evidence a
citation then opens to the wrong place. So a Word document parses as ONE page
(`ParsedPage(number=1, ...)`) with every paragraph and table in document
order — honest about what the format can support, rather than a page number
that LOOKS precise and is not.

TABLES: DOCX GETS THEM WHOLE; PDF DOES NOT, YET. `python-docx` walks
`document.element.body` in document order, so a `w:tbl` element becomes a
`ParsedTable` at exactly the position it appeared, interleaved correctly with
the paragraphs around it. `pypdf` has no comparable table API — genuine PDF
table detection needs a layout-analysis library outside this slab's
dependency budget (`requirements/ai.txt`'s `pypdf`/`python-docx`/
`markdown-it-py` three, nothing heavier). A PDF page's `tables` therefore
stays empty; its extracted text still carries table ROWS, just not as a
separately-addressable block. A later adapter fitting a table-aware PDF
library extends this seat without changing the port.
"""

from __future__ import annotations

from io import BytesIO

import docx
import pypdf

from cinqflow.ports import port
from cinqflow.ports.document_parse import (
    SUPPORTED_MEDIA_TYPES,
    DocumentParseError,
    ParsedDocument,
    ParsedPage,
    ParsedTable,
)

__all__ = ["FileDocumentParser"]


def _docx_row_text(row: object) -> tuple[str, ...]:
    """One `<w:tr>` element's cells, as plain text — `.iter()`'s document-order
    walk over each cell's `<w:t>` runs, which is what keeps a cell's text in
    reading order even when Word split it across several runs for formatting."""
    return tuple(
        "".join(node.text or "" for node in cell.iter() if node.tag.endswith("}t")).strip()
        for cell in row.iterchildren()  # type: ignore[attr-defined]
        if cell.tag.rsplit("}", 1)[-1] == "tc"
    )


@port("document_parse", "local-pypdf-docx")
class FileDocumentParser:
    def parse(self, content: bytes, *, media_type: str, filename: str = "") -> ParsedDocument:
        name = filename or "(unnamed)"
        if media_type not in SUPPORTED_MEDIA_TYPES:
            raise DocumentParseError(
                f"{name}: {media_type!r} is not a media type this adapter parses "
                f"({', '.join(sorted(SUPPORTED_MEDIA_TYPES))})"
            )
        if not content:
            raise DocumentParseError(f"{name}: no bytes to parse")

        if media_type == "application/pdf":
            return self._parse_pdf(content, name=name)
        if media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._parse_docx(content, name=name)
        # text/plain, text/markdown, text/csv — one page, decoded, no layout to lose.
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError as failure:
            raise DocumentParseError(f"{name}: not valid UTF-8 text — {failure}") from None
        if not text.strip():
            raise DocumentParseError(f"{name}: empty document")
        return ParsedDocument(media_type=media_type, pages=(ParsedPage(number=1, text=text),))

    def _parse_pdf(self, content: bytes, *, name: str) -> ParsedDocument:
        try:
            reader = pypdf.PdfReader(BytesIO(content))
        except Exception as failure:  # pypdf raises several distinct types on corrupt input
            raise DocumentParseError(f"{name}: not a readable PDF — {failure}") from None
        if reader.is_encrypted:
            raise DocumentParseError(f"{name}: encrypted PDFs are refused, not force-opened")

        pages: list[ParsedPage] = []
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(ParsedPage(number=index, text=text))
        if not pages:
            raise DocumentParseError(f"{name}: no extractable text on any page")
        return ParsedDocument(media_type="application/pdf", pages=tuple(pages))

    def _parse_docx(self, content: bytes, *, name: str) -> ParsedDocument:
        try:
            document = docx.Document(BytesIO(content))
        except Exception as failure:
            raise DocumentParseError(f"{name}: not a readable .docx — {failure}") from None

        blocks: list[str] = []
        tables: list[ParsedTable] = []
        for element in document.element.body.iterchildren():
            tag = element.tag.rsplit("}", 1)[-1]
            if tag == "p":
                text = "".join(
                    node.text or "" for node in element.iter() if node.tag.endswith("}t")
                )
                if text.strip():
                    blocks.append(text.strip())
            elif tag == "tbl":
                rows = tuple(
                    _docx_row_text(row)
                    for row in element.iterchildren()
                    if row.tag.rsplit("}", 1)[-1] == "tr"
                )
                if rows:
                    table = ParsedTable(rows=rows)
                    tables.append(table)
                    blocks.append(table.as_text())

        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if not blocks:
            raise DocumentParseError(f"{name}: no paragraphs or tables found")
        return ParsedDocument(
            media_type=media_type,
            pages=(ParsedPage(number=1, text="\n\n".join(blocks), tables=tuple(tables)),),
        )
