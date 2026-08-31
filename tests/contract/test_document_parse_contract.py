"""ONE suite, every parser. The `document_parse` pin's contract.

    "Parse layout-aware (tables kept whole, page anchors preserved)."
    — CF-V1-E16-04, acceptance criteria

The mock and the real adapter must agree on `text/plain` and `text/csv` —
those two media types have no layout to lose, so there is nothing for the
real adapter to do that the mock's honest default does not already do. PDF
and DOCX exercise the real adapter alone: the mock cannot fabricate genuine
page/table structure, and scripting a fake one would test the scripting
mechanism, not the parser.
"""

from __future__ import annotations

from collections.abc import Callable
from io import BytesIO

import docx
import pytest
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from cinqflow.adapters.local.file_document_parse import FileDocumentParser
from cinqflow.adapters.mock.document_parse import ScriptedDocumentParser
from cinqflow.ports.document_parse import (
    DocumentParseError,
    DocumentParsePort,
    ParsedDocument,
    ParsedPage,
)

pytestmark = pytest.mark.contract

PARSERS: dict[str, Callable[[], DocumentParsePort]] = {
    "mock": ScriptedDocumentParser,
    "local-pypdf-docx": FileDocumentParser,
}


@pytest.fixture(params=sorted(PARSERS))
def parser(request: pytest.FixtureRequest) -> DocumentParsePort:
    return PARSERS[request.param]()


# ── both adapters agree on the media types with no layout to lose ───────────


def test_plain_text_becomes_one_page_verbatim(parser: DocumentParsePort) -> None:
    content = b"BG-004: Member Identifier. The primary key on the enrollment file.\n"
    parsed = parser.parse(content, media_type="text/plain", filename="note.txt")
    assert parsed.page_count == 1
    assert parsed.pages[0].number == 1
    assert "BG-004" in parsed.pages[0].text


def test_csv_becomes_one_page_verbatim(parser: DocumentParsePort) -> None:
    content = b"Column,Format\nMBR_DOB,CCYYMMDD\n"
    parsed = parser.parse(content, media_type="text/csv", filename="spec.csv")
    assert parsed.page_count == 1
    assert "MBR_DOB" in parsed.pages[0].text


def test_empty_content_is_refused_not_returned_as_an_empty_document(
    parser: DocumentParsePort,
) -> None:
    with pytest.raises(DocumentParseError):
        parser.parse(b"", media_type="text/plain", filename="empty.txt")


def test_whitespace_only_content_is_refused(parser: DocumentParsePort) -> None:
    with pytest.raises(DocumentParseError):
        parser.parse(b"   \n\n  ", media_type="text/plain", filename="blank.txt")


# ── the real adapter, on real PDF and DOCX bytes ─────────────────────────────


def _one_page_pdf(text: str) -> bytes:
    """A minimal one-page PDF with real drawn text.

    `pypdf.PdfWriter` can compose a page but has no text-drawing API of its
    own (it manipulates PDFs; it does not generate their content streams) —
    so this fixture writes a minimal `BT ... Tj ET` content stream directly,
    which is enough for `extract_text()` to return `text` and nothing more.
    """
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    page = writer.pages[0]
    stream = DecodedStreamObject()
    stream.set_data(f"BT /F1 12 Tf 20 100 Td ({text}) Tj ET".encode())
    page[NameObject("/Contents")] = writer._add_object(stream)
    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {
                    NameObject("/F1"): DictionaryObject(
                        {
                            NameObject("/Type"): NameObject("/Font"),
                            NameObject("/Subtype"): NameObject("/Type1"),
                            NameObject("/BaseFont"): NameObject("/Helvetica"),
                        }
                    )
                }
            )
        }
    )
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def test_a_real_pdf_yields_page_anchored_text() -> None:
    content = _one_page_pdf("MBR_DOB is CCYYMMDD")
    parsed = FileDocumentParser().parse(content, media_type="application/pdf", filename="guide.pdf")
    assert parsed.media_type == "application/pdf"
    assert parsed.page_count == 1
    assert parsed.pages[0].number == 1
    assert "MBR_DOB" in parsed.pages[0].text


def test_a_corrupt_pdf_is_refused_by_name() -> None:
    with pytest.raises(DocumentParseError, match="not a readable PDF"):
        FileDocumentParser().parse(
            b"not a pdf at all", media_type="application/pdf", filename="x.pdf"
        )


def _docx_with_table() -> bytes:
    document = docx.Document()
    document.add_paragraph("The companion guide's field mapping:")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Column"
    table.cell(0, 1).text = "Format"
    table.cell(1, 0).text = "MBR_DOB"
    table.cell(1, 1).text = "CCYYMMDD"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


DOCX_MEDIA = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def test_a_real_docx_keeps_its_table_whole() -> None:
    parsed = FileDocumentParser().parse(
        _docx_with_table(), media_type=DOCX_MEDIA, filename="spec.docx"
    )
    assert parsed.page_count == 1
    page = parsed.pages[0]
    assert "companion guide" in page.text
    assert len(page.tables) == 1
    assert page.tables[0].rows == (
        ("Column", "Format"),
        ("MBR_DOB", "CCYYMMDD"),
    )
    # The table's own text stays intact — a chunker never splits mid-row.
    assert "MBR_DOB | CCYYMMDD" in page.tables[0].as_text()


def test_a_corrupt_docx_is_refused_by_name() -> None:
    with pytest.raises(DocumentParseError, match=r"not a readable \.docx"):
        FileDocumentParser().parse(b"not a docx", media_type=DOCX_MEDIA, filename="x.docx")


def test_an_unsupported_media_type_is_refused_by_name() -> None:
    with pytest.raises(DocumentParseError, match="image/png"):
        FileDocumentParser().parse(b"\x89PNG", media_type="image/png", filename="x.png")


# ── the mock is scriptable for the two formats it cannot fabricate ──────────


def test_the_mock_can_be_scripted_for_pdf_and_docx() -> None:
    mock = ScriptedDocumentParser()
    content = b"scripted pdf bytes"
    mock.script(
        content,
        media_type="application/pdf",
        result=ParsedDocument(
            media_type="application/pdf", pages=(ParsedPage(number=14, text="MBR_DOB: CCYYMMDD"),)
        ),
    )
    parsed = mock.parse(content, media_type="application/pdf", filename="guide.pdf")
    assert parsed.pages[0].number == 14


def test_the_mock_can_be_scripted_to_refuse() -> None:
    mock = ScriptedDocumentParser()
    mock.refuse(b"bad", media_type="application/pdf", reason="scripted refusal")
    with pytest.raises(DocumentParseError, match="scripted refusal"):
        mock.parse(b"bad", media_type="application/pdf", filename="x.pdf")


def test_an_unscripted_pdf_on_the_mock_is_refused_not_guessed() -> None:
    with pytest.raises(DocumentParseError, match="no scripted result"):
        ScriptedDocumentParser().parse(b"anything", media_type="application/pdf", filename="x.pdf")
